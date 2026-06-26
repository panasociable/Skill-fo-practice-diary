"""Собирает template.docx с {{плейсхолдерами}} из исходного дневника ПГТУ."""
from docx import Document
import re

SRC = "tools/template_src.docx"
OUT = "skills/dnevnik-praktiki/assets/template.docx"

def set_text(p, new):
    if not p.runs:
        p.add_run(new); return
    p.runs[0].text = new
    for r in p.runs[1:]:
        r.text = ""

SUBS = {
    "производственная": "{{VID}}",
    "организационно-управленческая": "{{TIP}}",
    "Петров Петр Петрович": "{{FIO}}",
    "кафедра менеджмента и бизнеса": "{{KAFEDRA}}",
    "27.03.05 Инноватика": "{{SPEC}}",
    "очная": "{{FORMA}}",
    "ООО «Принтекс»": "{{MESTO}}",
    "директор – Мокеичев А.Н.": "{{INSTR_ORG}}",
    "Написать содержание вашего отчета": "{{ZADANIE}}",
}
FULL = {
    "Курс__2____ Группа__ИНВ-21______":
        "Курс__{{KURS}}____ Группа__{{GRUPPA}}______",
    "Сроки практики с «_29_»_июня____2026г. по «_12___»_июля___2026 г.":
        "Сроки практики с «_{{D1}}_»_{{M1}}____{{Y1}}г. по «_{{D2}}_»_{{M2}}___{{Y2}} г.",
    "Приказ по университету/филиалу № _______от «_____»_____________2026г.":
        "Приказ по университету/филиалу № _{{PRIKAZ_N}}__от «_{{PRIKAZ_D}}_»_{{PRIKAZ_M}}__{{PRIKAZ_Y}}г.",
    "Договор о проведении практической подготовки №_______ от «____»__________20___г.":
        "Договор о проведении практической подготовки №_{{DOG_N}}_ от «_{{DOG_D}}_»_{{DOG_M}}_{{DOG_Y}}г.",
    "2026г.": "{{YEAR}}г.",
}
PERIODS = ["29.06.-30.06","01.07-02.07","03.07-05.07","05.07-06.07",
           "07.07-08.07","09.07-10.07","11.07-11.07","11.07-10.07"]

doc = Document(SRC)
for p in doc.paragraphs:
    t = p.text
    if t in FULL:
        set_text(p, FULL[t]); continue
    new = t
    for k, v in SUBS.items():
        if k in new: new = new.replace(k, v)
    if new != t: set_text(p, new)

for i, p in enumerate(doc.paragraphs):
    if i == 41 and p.text.strip() and set(p.text.strip()) <= set("_ "):
        set_text(p, "________________{{INSTR_VUZ}}________________________________________")

tbl = doc.tables[1]
for r in tbl.rows:
    for c in r.cells:
        for k_i, per in enumerate(PERIODS, 1):
            if per in c.text:
                for p in c.paragraphs:
                    if per in p.text:
                        set_text(p, p.text.replace(per, "{{P%d}}" % k_i))

doc.save(OUT)
d2 = Document(OUT)
found = set()
for p in d2.paragraphs:
    found.update(re.findall(r"\{\{[A-Z0-9_]+\}\}", p.text))
for t in d2.tables:
    for r in t.rows:
        for c in r.cells:
            found.update(re.findall(r"\{\{[A-Z0-9_]+\}\}", c.text))
print("OK ->", OUT)
print("Плейсхолдеры:", sorted(found))
