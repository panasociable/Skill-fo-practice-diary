#!/usr/bin/env python3
"""Заполняет дневник практики ПГТУ из JSON с данными студента.

Запуск:
    python3 fill_diary.py data.json "Дневник_Иванов.docx"

Поля JSON (см. assets/data_example.json). Обязательные: fio, kafedra, spec,
kurs, gruppa, mesto, start_date, end_date, zadanie.
Даты — в формате ГГГГ-ММ-ДД. План работ (8 пунктов) и его даты
рассчитываются автоматически по срокам практики.

При неверных данных скрипт печатает понятную ошибку и завершается с кодом 1,
а не сырым traceback.
"""
import sys, json, os
from datetime import date, timedelta
from docx import Document

MONTHS_GEN = ["", "января","февраля","марта","апреля","мая","июня",
              "июля","августа","сентября","октября","ноября","декабря"]

REQUIRED = ["fio","kafedra","spec","kurs","gruppa","mesto",
            "start_date","end_date","zadanie"]

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(HERE, "..", "assets", "template.docx")


class InputError(Exception):
    """Понятная ошибка во входных данных (без traceback для пользователя)."""


def die(msg):
    print("Ошибка: " + msg, file=sys.stderr)
    sys.exit(1)


def parse_date(s, field):
    if not isinstance(s, str):
        raise InputError("поле «%s» должно быть строкой с датой в формате ГГГГ-ММ-ДД" % field)
    try:
        y, m, d = s.split("-")
        return date(int(y), int(m), int(d))
    except (ValueError, AttributeError):
        raise InputError(
            "поле «%s» = %r — неверный формат даты. Нужен ГГГГ-ММ-ДД, например 2026-06-29" % (field, s)
        )


def validate(data):
    if not isinstance(data, dict):
        raise InputError("JSON должен быть объектом с полями, а не %s" % type(data).__name__)
    missing = [f for f in REQUIRED if not str(data.get(f, "")).strip()]
    if missing:
        raise InputError("не заполнены обязательные поля: " + ", ".join(missing))

    s = parse_date(data["start_date"], "start_date")
    e = parse_date(data["end_date"], "end_date")
    if e < s:
        raise InputError(
            "end_date (%s) раньше start_date (%s) — проверьте сроки практики"
            % (data["end_date"], data["start_date"])
        )
    days = (e - s).days + 1
    if days < 8:
        # не критично, но план из 8 пунктов будет с повторяющимися датами
        print("Предупреждение: срок практики всего %d дн., "
              "даты в плане из 8 пунктов будут перекрываться." % days, file=sys.stderr)

    # курс — число
    try:
        int(str(data["kurs"]))
    except ValueError:
        raise InputError("поле «kurs» должно быть числом, получено %r" % data["kurs"])

    # необязательные даты, если заданы
    for f in ("prikaz_date", "dog_date"):
        if str(data.get(f, "")).strip():
            parse_date(data[f], f)
    return s, e


def split_periods(start, end, n=8):
    """Делит срок практики на n этапов, формат 'ДД.ММ-ДД.ММ'."""
    total = (end - start).days + 1
    if total < n:
        total = n
    out = []
    for i in range(n):
        a = start + timedelta(days=(total * i) // n)
        b = start + timedelta(days=(total * (i + 1)) // n - 1)
        if b < a: b = a
        if b > end: b = end
        out.append("%02d.%02d-%02d.%02d" % (a.day, a.month, b.day, b.month))
    return out


def set_text(p, new):
    if not p.runs:
        p.add_run(new); return
    p.runs[0].text = new
    for r in p.runs[1:]:
        r.text = ""


def replace_all(doc, mapping):
    def fix(p):
        if "{{" not in p.text: return
        t = p.text
        for k, v in mapping.items():
            t = t.replace(k, v)
        set_text(p, t)
    for p in doc.paragraphs: fix(p)
    for tb in doc.tables:
        for r in tb.rows:
            for c in r.cells:
                for p in c.paragraphs: fix(p)


def build_mapping(data, s, e):
    periods = split_periods(s, e)
    g = data.get
    m = {
        "{{VID}}":    g("vid", "производственная"),
        "{{TIP}}":    g("tip", ""),
        "{{FIO}}":    data["fio"],
        "{{KAFEDRA}}": data["kafedra"],
        "{{SPEC}}":   data["spec"],
        "{{FORMA}}":  g("forma", "очная"),
        "{{KURS}}":   str(data["kurs"]),
        "{{GRUPPA}}": data["gruppa"],
        "{{MESTO}}":  data["mesto"],
        "{{D1}}": str(s.day), "{{M1}}": MONTHS_GEN[s.month], "{{Y1}}": str(s.year),
        "{{D2}}": str(e.day), "{{M2}}": MONTHS_GEN[e.month], "{{Y2}}": str(e.year),
        "{{YEAR}}": str(s.year),
        "{{INSTR_VUZ}}": g("instr_vuz", ""),
        "{{INSTR_ORG}}": g("instr_org", ""),
        "{{ZADANIE}}": data["zadanie"],
    }

    def datebits(key, pre):
        v = data.get(key)
        if str(v or "").strip():
            d = parse_date(v, key)
            m["{{%s_D}}" % pre] = str(d.day)
            m["{{%s_M}}" % pre] = MONTHS_GEN[d.month]
            m["{{%s_Y}}" % pre] = str(d.year)
        else:
            m["{{%s_D}}" % pre] = "___"
            m["{{%s_M}}" % pre] = "________"
            m["{{%s_Y}}" % pre] = "20__"

    m["{{PRIKAZ_N}}"] = g("prikaz_n", "______")
    m["{{DOG_N}}"]    = g("dog_n", "______")
    datebits("prikaz_date", "PRIKAZ")
    datebits("dog_date", "DOG")
    for i, p in enumerate(periods, 1):
        m["{{P%d}}" % i] = p
    return m


def generate(data, out_path):
    """Генерирует .docx. Может бросить InputError при плохих данных."""
    s, e = validate(data)
    if not os.path.exists(TEMPLATE):
        raise InputError("не найден шаблон: %s" % TEMPLATE)
    doc = Document(TEMPLATE)
    replace_all(doc, build_mapping(data, s, e))
    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 2:
        die("укажите путь к JSON: python3 fill_diary.py data.json [output.docx]")
    json_path = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else "Дневник_практики.docx"
    if not os.path.exists(json_path):
        die("файл не найден: %s" % json_path)
    try:
        with open(json_path, encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError as ex:
        die("некорректный JSON в %s: %s" % (json_path, ex))
    try:
        generate(data, out)
    except InputError as ex:
        die(str(ex))
    print("Готово:", out)


if __name__ == "__main__":
    main()
