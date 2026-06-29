"""Автотесты генератора дневника. Запуск:

    pytest                       # если установлен pytest
    python3 tests/test_fill_diary.py   # standalone, без pytest
"""
import os, re, sys, json, tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SCRIPTS = os.path.join(ROOT, "skills", "dnevnik-praktiki", "scripts")
ASSETS = os.path.join(ROOT, "skills", "dnevnik-praktiki", "assets")
sys.path.insert(0, SCRIPTS)

import fill_diary  # noqa: E402
from docx import Document  # noqa: E402

PLACEHOLDER = re.compile(r"\{\{.*?\}\}")


def _all_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(c.text)
    return "\n".join(parts)


def test_example_has_no_placeholders():
    """Документ из примера генерируется и не содержит {{...}}."""
    data = json.load(open(os.path.join(ASSETS, "data_example.json"), encoding="utf-8"))
    out = os.path.join(tempfile.gettempdir(), "test_dnevnik.docx")
    fill_diary.generate(data, out)
    text = _all_text(out)
    left = PLACEHOLDER.findall(text)
    assert not left, "Остались незаполненные плейсхолдеры: %s" % set(left)
    assert data["fio"] in text
    assert data["mesto"] in text


def test_missing_required_field_raises():
    """Отсутствие обязательного поля даёт понятную ошибку, а не traceback."""
    data = {"fio": "Тест"}
    try:
        fill_diary.generate(data, "/tmp/x.docx")
    except fill_diary.InputError as ex:
        assert "обязательные поля" in str(ex)
    else:
        raise AssertionError("ожидалась InputError")


def test_bad_date_raises():
    data = json.load(open(os.path.join(ASSETS, "data_example.json"), encoding="utf-8"))
    data["start_date"] = "29.06.2026"  # неверный формат
    try:
        fill_diary.generate(data, "/tmp/x.docx")
    except fill_diary.InputError as ex:
        assert "формат даты" in str(ex)
    else:
        raise AssertionError("ожидалась InputError")


def test_end_before_start_raises():
    data = json.load(open(os.path.join(ASSETS, "data_example.json"), encoding="utf-8"))
    data["start_date"], data["end_date"] = "2026-07-12", "2026-06-29"
    try:
        fill_diary.generate(data, "/tmp/x.docx")
    except fill_diary.InputError as ex:
        assert "раньше" in str(ex)
    else:
        raise AssertionError("ожидалась InputError")


if __name__ == "__main__":
    fails = 0
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            try:
                fn(); print("OK   ", name)
            except AssertionError as e:
                fails += 1; print("FAIL ", name, "-", e)
    sys.exit(1 if fails else 0)
